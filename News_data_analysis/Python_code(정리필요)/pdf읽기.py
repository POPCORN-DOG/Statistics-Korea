import pandas as pd
import numpy as np
# pdf파일을 text문서로 가져오기
'''연습(실패)
#PyPDF2 패키지 임포트 하기
import PyPDF2

#파일 열기

file = open('C:/Users/USER/Desktop/디지털시대의 고용안전망-4_web.pdf','rb')

filereader = PyPDF2.PdfFileReader(file)
print(filereader.getPage(50))
print(filereader.numPages)

#첫번째 페이지 정보 가져오기
page0 = filereader.getPage(100)

filereader.read()
#text로 가져오기
text = page0.extractText()
text.encode('utf-16').decode('utf-16')
text.decode('ascii')

text.decode('Identity-H')

page0.Encoding()

from tika import parser

file = parser.from_file('C:/Users/USER/Desktop/디지털시대의 고용안전망-4_web.pdf')
file1 = file['content'].splitlines()

while len(file1[0]) == 0:
    file1.pop(0)
file1[10004]

'''

#doc파일 불러오기

from docx import Document
file = Document('C:/Users/USER/Desktop/고용보고서/2019년도 사업체패널 워킹페이퍼 시리즈(web).docx')
len(file.paragraphs)
file.text()

#파일 한번에 불러와서 저장하기
import string
import glob
import os
from tqdm import tqdm
import re
flist = os.path.join('C:/Users/USER/Desktop/고용보고서', '*.docx') #끝에 합본.csv로 끝나는 것들 불러오기
file_list = glob.glob(flist)  #파일 list 불러오기

for i in tqdm(file_list):
    file = Document(i)
    a = []
    for x, paragraph in enumerate(file.paragraphs):
        a.append(paragraph.text)
        #print(str(x)+ '|' + str(len(paragraph.text)) + " : " + paragraph.text)
    fw = open(i[:-5] + '한글만.txt', 'w', encoding='utf-8')
    fw.write(' '.join(re.compile('[|가-힣]+').findall(' '.join(a))))
    fw.close()

' '.join(re.compile('[|가-힣]+').findall(' '.join(a)))

'''혹시 읽기 안되면
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.detach(),encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.detach(),encoding='utf-8')'''

###연관성(네트워크) 분석

##보고서 다시 불러오기
flist = os.path.join('C:/Users/USER/Desktop/고용보고서', '*.docx') #끝에 합본.csv로 끝나는 것들 불러오기
file_list = glob.glob(flist)  #파일 list 불러오기

a = []
for i in tqdm(file_list):
    file = Document(i)
    for x, paragraph in enumerate(file.paragraphs):
        a.append(paragraph.text)

a1 = [' '.join(re.compile('[|가-힣]+').findall(''.join(i))) for i in a] #한글만 가져오기
a1 = [i for i in a1 if len(i) > 0] #내용없는 단락 없애기 전처리 하려면 0을 50이나 100으로 키우기

##형태소 분석하기
from ckonlpy.tag import Twitter
#단어 데이터 불러오기
work_word_data = pd.read_excel('C:/Users/USER/Desktop/고용관련단어.xlsx')
#단어 리스트 만들기
work_word_list = [i for i in work_word_data['선정용어']]
work_word_list

#밑줄 단어 만들기
work_word_list2 = []
for i in  work_word_list:
    work_word_list2.append(i.replace(' ','_'))

#띄어쓰기 없는단어 만들기
work_word_list3 = []
for i in  work_word_list:
    work_word_list3.append(i.replace(' ',''))

#단어 리스트 형태소 분석기에 넣기
twi = Twitter()
for i in range(len(work_word_list2)):
    twi.add_dictionary(work_word_list2[i],'Noun')
    twi.add_dictionary(work_word_list[i],'Noun')
    twi.add_dictionary(work_word_list3[i],'Noun')

#통합어 딕셔너리 만들기
replace = {}
for i in range(len(work_word_list2)):
    replace[work_word_list2[i]] = work_word_list[i]
    replace[work_word_list3[i]] = work_word_list[i]
#기타 통합어 추가
replace['특고'] = '특수형태근로종사자'
replace['특고종사자'] = '특수형태근로종사자'
replace['특수고용직'] = '특수형태근로종사자'
replace['맞벌이'] = '맞벌이 가구'

#형태소 분석
from ckonlpy.tag import Postprocessor
postprocessor = Postprocessor(base_tagger=twi, replace=replace, passtags={'Noun'})
under_bar_text = [m.replace(' ','_') for m in a1] #본문에도 띄어쓰기를 _로 바꿔야 하는데 못바꿈 2.0버젼에 적용
#형태소 분석 하기
words = [[j[i][0] for i in range(len(j))] for j in [postprocessor.pos(i) for i in tqdm(under_bar_text)]]

#통계 선정용어만 가져오기
j = 0
w =[]
words2 = []
for i in tqdm(words):
    for k in work_word_data['선정용어']:
        if k in i:
            words2.append(k)
    w.append(words2)
    words2 = []
    j += 1 #불용어 제외
w2 = [i for i in w if len(i) > 1]

#연관어 분석 해보기
from apyori import apriori
import networkx as nx
from mlxtend.frequent_patterns import association_rules
from mlxtend.preprocessing import TransactionEncoder

dataset = w2

#신뢰도 보기
te = TransactionEncoder()
te_ary = te.fit(dataset).transform(dataset)
df = pd.DataFrame(te_ary, columns=te.columns_)
frequent_itemsets = apriori(dataset, min_support=0.001, use_colnames = True)
a = pd.DataFrame(frequent_itemsets, columns=['itemsets','support','ordered_statistics'])
confidence = association_rules(a, metric='confidence',min_threshold=0.3)
#지지도

result = apriori(dataset, min_support=0.001)
df = pd.DataFrame(result)
df['length'] = df['items'].apply(lambda x: len(x))
df = df[(df['length'] == 2) & (df['support'] >= 0.001)].sort_values(by = 'support', ascending= False)

#그래프 정의
g = nx.Graph()
ar = (df['items'])
g.add_edges_from(ar)

#페이지 랭크
pr = nx.pagerank(g)
nsize = np.array([v for v in pr.values()])
nsize = 2000 * (nsize - min(nsize)) / (max(nsize) - min(nsize))

#레이아웃
pos = nx.planar_layout(g)
pos = nx.shell_layout(g)
pos = nx.spring_layout(g)
pos = nx.kamada_kawai_layout(g)


#글씨체
from matplotlib import font_manager, rc
import matplotlib
import matplotlib.pyplot as plt
font_name = font_manager.FontProperties(fname='C:/Windows/Fonts/H2HDRM.TTF').get_name()
rc('font', family=font_name)
matplotlib.rcParams['axes.unicode_minus'] = False

plt.figure(figsize=(16,12)); plt.axis('off')
nx.draw_networkx(g, pos=pos, font_family = font_name, node_size = nsize,
                 alpha = 0.7, edge_color = '.3',node_color = list(pr.values()))# , cmap = plt.cm.rainbow ,node_color = pr.values())


#######너무 커서 안 읽힘 나누기
'''
flist = os.path.join('C:/Users/USER/Desktop/고용보고서', '*.docx') #끝에 합본.csv로 끝나는 것들 불러오기
file_list = glob.glob(flist)  #파일 list 불러오기

a = []
for i in tqdm(file_list):
    file = Document(i)
    for x, paragraph in enumerate(file.paragraphs):
        a.append(paragraph.text)

a1 = [' '.join(re.compile('[|가-힣]+').findall(''.join(i))) for i in a] #한글만 가져오기
a1 = [i for i in a1 if len(i) > 0] #내용없는 단락 없애기 전처리 하려면 0을 50이나 100으로 키우기

int((len(a1)/60) * 1)
a1[0:int((len(a1)/60) * 1)]

for i in range(0,120):
    b1 =  a1[int((len(a1) / 120) * i):int((len(a1) / 120) * (i+1))]
    fw = open('C:/Users/USER/Desktop/고용보고서/나누기/' + str(i) + '.txt', 'w', encoding='utf-8')
    fw.write(' '.join(b1))
    fw.close()'''




